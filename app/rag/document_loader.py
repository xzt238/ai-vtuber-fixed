"""
文档加载器

支持多种文档格式的加载和解析。
"""

import os
import hashlib
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass

from . import Document
import logging

logger = logging.getLogger(__name__)


class DocumentLoader:
    """文档加载器"""
    
    def __init__(self, config: Dict[str, Any] = None):
        self.config = config or {}
        self.supported_formats = self.config.get("supported_formats", ["pdf", "txt", "md", "docx"])
    
    def load(self, file_path: str) -> Optional[Document]:
        """加载文档"""
        try:
            file_path = Path(file_path)
            
            # 检查文件是否存在
            if not file_path.exists():
                logger.info(f" 文件不存在: {file_path}")
                return None
            
            # 检查文件格式
            file_type = file_path.suffix.lower().lstrip('.')
            if file_type not in self.supported_formats:
                logger.info(f" 不支持的文件格式: {file_type}")
                return None
            
            # 生成文档ID
            doc_id = self._generate_doc_id(file_path)
            
            # 根据文件类型加载内容
            content = self._load_content(file_path, file_type)
            if content is None:
                return None
            
            # 提取元数据
            metadata = self._extract_metadata(file_path, file_type)
            
            # 创建文档对象
            document = Document(
                id=doc_id,
                file_path=str(file_path),
                file_name=file_path.name,
                file_type=file_type,
                content=content,
                metadata=metadata
            )
            
            logger.info(f" 文档加载成功: {file_path.name}")
            logger.info(f" 文件类型: {file_type}, 内容长度: {len(content)} 字符")
            
            return document
            
        except Exception as e:
            logger.info(f" 文档加载失败: {e}")
            return None
    
    def load_batch(self, file_paths: List[str]) -> List[Document]:
        """批量加载文档"""
        documents = []
        for file_path in file_paths:
            doc = self.load(file_path)
            if doc is not None:
                documents.append(doc)
        return documents
    
    def _generate_doc_id(self, file_path: Path) -> str:
        """生成文档ID"""
        # 使用文件路径和大小生成唯一ID
        stat = file_path.stat()
        content = f"{file_path}:{stat.st_size}:{stat.st_mtime}"
        return hashlib.md5(content.encode()).hexdigest()
    
    def _load_content(self, file_path: Path, file_type: str) -> Optional[str]:
        """加载文档内容"""
        try:
            if file_type == "txt":
                return self._load_txt(file_path)
            elif file_type == "md":
                return self._load_markdown(file_path)
            elif file_type == "pdf":
                return self._load_pdf(file_path)
            elif file_type == "docx":
                return self._load_docx(file_path)
            else:
                logger.info(f" 未实现的文件类型: {file_type}")
                return None
        except Exception as e:
            logger.info(f" 内容加载失败: {e}")
            return None
    
    def _load_txt(self, file_path: Path) -> Optional[str]:
        """加载TXT文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # 尝试其他编码
            try:
                with open(file_path, 'r', encoding='gbk') as f:
                    return f.read()
            except (UnicodeDecodeError, IOError, OSError):
                return None
    
    def _load_markdown(self, file_path: Path) -> Optional[str]:
        """加载Markdown文件"""
        return self._load_txt(file_path)
    
    def _load_pdf(self, file_path: Path) -> Optional[str]:
        """加载PDF文件"""
        try:
            # 尝试导入PyPDF2
            try:
                import PyPDF2
            except ImportError:
                logger.info(" PyPDF2 未安装，无法加载PDF文件")
                logger.info(" 请运行: pip install PyPDF2")
                return None
            
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                text_parts = []
                
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                
                return '\n\n'.join(text_parts)
                
        except Exception as e:
            logger.info(f" PDF加载失败: {e}")
            return None
    
    def _load_docx(self, file_path: Path) -> Optional[str]:
        """加载DOCX文件"""
        try:
            # 尝试导入python-docx
            try:
                import docx
            except ImportError:
                logger.info(" python-docx 未安装，无法加载DOCX文件")
                logger.info(" 请运行: pip install python-docx")
                return None
            
            doc = docx.Document(file_path)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    text_parts.append(paragraph.text)
            
            return '\n\n'.join(text_parts)
            
        except Exception as e:
            logger.info(f" DOCX加载失败: {e}")
            return None
    
    def _extract_metadata(self, file_path: Path, file_type: str) -> Dict[str, Any]:
        """提取文档元数据"""
        stat = file_path.stat()
        
        metadata = {
            "file_name": file_path.name,
            "file_type": file_type,
            "file_size": stat.st_size,
            "created_time": stat.st_ctime,
            "modified_time": stat.st_mtime,
            "file_path": str(file_path),
        }
        
        # 根据文件类型提取特定元数据
        if file_type == "pdf":
            metadata.update(self._extract_pdf_metadata(file_path))
        elif file_type == "docx":
            metadata.update(self._extract_docx_metadata(file_path))
        
        return metadata
    
    def _extract_pdf_metadata(self, file_path: Path) -> Dict[str, Any]:
        """提取PDF元数据"""
        try:
            import PyPDF2
            
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                info = reader.metadata
                
                if info:
                    return {
                        "title": info.title,
                        "author": info.author,
                        "subject": info.subject,
                        "creator": info.creator,
                        "pages": len(reader.pages),
                    }
        except Exception:
            pass
        
        return {}
    
    def _extract_docx_metadata(self, file_path: Path) -> Dict[str, Any]:
        """提取DOCX元数据"""
        try:
            import docx
            
            doc = docx.Document(file_path)
            props = doc.core_properties
            
            return {
                "title": props.title,
                "author": props.author,
                "subject": props.subject,
                "keywords": props.keywords,
            }
        except Exception:
            pass
        
        return {}